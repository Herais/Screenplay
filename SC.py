# -*- coding: utf-8 -*-
"""
Created on Mon Feb 17 00:31:27 2020

@author: VXhpUS
"""

import numpy as np
import pandas as pd
import re
from docx import Document
import time
from xml.dom import minidom
import lxml.etree as et
import codecs


## For Translate
import hashlib
from random import randint
import http.client
import urllib
import json

#%% Class Screenplay
class Screenplay(object):
    
    def __init__(self, sc: str = None):   

        super(Screenplay, self).__init__()
        self.sc = sc
        self.elements = Elements()
        self.reformat = Reformat.to_OpenXML()
        self.element2basestyle = {
                    'text':'Normal Text',
                    'h': 'Scene Heading',
                    'a': 'Action',
                    'c': 'Character',
                    'p': 'Parenthetical',
                    'd': 'Dialogue',
                    't': 'Transition',
                    's': 'Shot',
                    'uc': 'Unspoken Character'
                    }
        self.translate = Translate()
        
        # Scene heading element
    
    def open(self, scpath: str):
        pass
    
    def read_openxml(self, filepath: str) -> pd.DataFrame:
        pass
    
    def read_docx(self, filepath: str) -> pd.DataFrame:
        '''
        This function reads word docx into a screenplay structured pd.DataFrame, 
        oneline per row.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        sctype : str, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        script : pd.DataFrame 
            DESCRIPTION.

        '''
        document = Document(filepath)
        num_p = len(document.paragraphs)
        script = []
        for num in range(num_p):
            script.append(document.paragraphs[num].text)
        script = pd.Series(script).rename('pcontent')
        script = script[script != ''] # remove empty lines
        script = pd.DataFrame(script) # convert to DataFrame
        
        # Paragraph type
        '''
        h = scene heading
            - Setting (i.e. INT/EXT)
            - Parenthesis (i.e. (Flashback))
        d = dialog
           - Characters
           - Parenthesis
        t = text
        s = shot (i.e. POV)
        t = transition (i.e. CUT TO, DISSOLVE, FADE, CROSS FADE)
            
        '''
        script['ptype'] = np.nan 
        
        # Scene Heading elements
        script['h_number'] = None
        script['h_inout'] = None
        script['h_title'] = None
        script['h_time'] = None
        script['h_parenthesis'] = None
        script['h_characters_in_scene'] = None
        
        # Dialog 
        script['d_character'] = None
        script['d_character_parenthesis'] = None
        script['d_dialog'] = None
        script['d_dialog_parenthesis'] = None
        return script
    
    def export_csv(self, filepath):
        self.sc.to_csv(filepath, encoding='utf-8-sig')
    
    
class Elements(object):
    
    def __init__(self):
        super(Elements, self).__init__()
        
    def identify_scene_heading(self, sc, pattern: str = '\d*[.\s]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'h'
        return sc
    
    def identify_dialog(self, sc, pattern: str = u'[：:]') -> pd.DataFrame:
        sc.loc[sc['pcontent'].apply(lambda x: x[0:10]).str.contains(pattern), 'ptype'] = 'd'
        return sc
    
    def identify_action(self, sc, pattern: str = None) -> pd.DataFrame:
        if pattern:
            sc.loc[sc['pcontent'].str.contains(pattern), 'ptype'] = 'a'
        else:
            sc['ptype'].fillna(value='a', inplace=True)
        return sc
    
    # Scene Heading Subelements
    def identify_scene_number(self, 
                              sc, 
                              pattern: str, 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_number'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_number'].fillna(method='ffill', inplace=True)
        return sc
    
    def identify_scene_inout(self, 
                              sc, 
                              pattern: str = r"[\w\S]*(.*?)[，,]", 
                              identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_inout'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0].str.lstrip(' ')
        #sc['h_inout'].fillna(method='ffill', inplace=True)
        return sc
        pass
        
    def identify_scene_title(self, 
                          sc, 
                          pattern: str = r'景[，,](\s*\w*.*)[，,]', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_title'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_title'].fillna(method='ffill', inplace=True)
        return sc

    def identify_scene_time(self, 
                          sc, 
                          pattern: str = r'[，,](\s*\w*.*$)', 
                          identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_scene_heading(sc, pattern)
        sc.loc[sc['ptype'] == 'h', 'h_time'] = sc.loc[sc['ptype'] == 'h', 'pcontent'].str.extract(pattern).iloc[:, 0]
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc
    
    # Dialog Subelements
    def identify_dialog_character(self,
                                  sc,
                                  pattern: str = r'(^.*$)[:：]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
        sc  = self.identify_dialog_character_parenthesis(sc)
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc

    def identify_dialog_character_parenthesis(self,
                                  sc: pd.DataFrame,
                                  pattern: str = r'[（(](.*)[）)]',
                                  identify_type_first: bool = False) -> pd.DataFrame:
        if identify_type_first:
            sc = self.identify_dialog(sc, pattern)
        sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.extract(pattern).iloc[:, 0]
        sc.loc[sc['ptype'] == 'd', 'd_character'] = sc.loc[sc['ptype'] == 'd', 'd_character'].str.replace(pattern, '')
        #sc['h_time'].fillna(method='ffill', inplace=True)
        return sc        
                        
    def identify_dialog_dialog(self,
                                 sc,
                                 pattern: str = r'[:：](.*)',
                                 identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'pcontent'].str.extract(pattern).iloc[:, 0]
       sc.loc[sc['ptype'] == 'd', 'd_dialog'] = sc.loc[sc['ptype'] == 'd', 'd_dialog'].apply(lambda x: x.lstrip(' ').rstrip(' '))
       sc  = self.identify_dialog_parenthetisis(sc)
       return sc
   
    def identify_dialog_parenthetisis(self,
                                      sc: pd.DataFrame,
                                      pattern: str = r'[（(](.*)[）)]',
                                      identify_type_first: bool = False) -> pd.DataFrame:
       if identify_type_first:
           sc = self.identify_dialog(sc, pattern)
       sc.loc[sc['ptype'] == 'd', 'd_dialog_parenthesis'] = sc.loc[sc['ptype'] == 'd', 'd_dialog']. \
           str.extractall(r'[（(](.*?)[）)]').reset_index().groupby('level_0')[0].agg(';'.join).str.split(';')


       #sc['h_time'].fillna(method='ffill', inplace=True)
       return sc
   
######################################################
class Reformat(object):
    
    def __init__(self):
        super(Reformat, self).__init__()
    
    class to_OpenXML(object):
        
        def __init__(self):
            pass
    
        def to_openxml(self, sc: pd.DataFrame, save: bool = False, file_path: str  = None) -> str:
            '''
            
    
            Parameters
            ----------
            sc : pd.DataFrame
                DESCRIPTION.
            save : str, optional
                DESCRIPTION. The default is None.
            file_path : str, optional
                DESCRIPTION. The default is None.
    
            Returns
            -------
            str
                DESCRIPTION.
    
            '''
            sc['formatted'] = None
    
            header_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
            header_document = '<document type="Open Screenplay Format document" version="30">\n'
            document_info = '\t<info>\n\t</info>\n'            
            header_format = '\t<styles>\n\t</styles>\n'
            header_paragraphs = '\t<paragraphs>\n'
            
            footer_paragraphs = '\t</paragraphs>\n'
            section_titlepage = '\t<titlepage>\n\t</titlepage>\n'
            section_lists = '\t<lists>\n\t</lists>\n'
            footer_document = '</document>\n'
                
            sc = self.heading(sc)
            sc = self.action(sc)
            sc = self.dialog(sc)
    
            formatted_output = header_xml \
                             + header_document \
                             + document_info \
                             + header_format \
                             + header_paragraphs \
                             + sc['formatted'].sum() \
                             + footer_paragraphs \
                             + section_titlepage \
                             + section_lists \
                             + footer_document \
                                 
            if save:
                with codecs.open(file_path, "w", 'utf-8-sig') as f:
                    print(formatted_output, file=f)
                f.close()
                
            return formatted_output
        
        
        def heading(self, sc: pd.DataFrame) -> pd.DataFrame:
            sc.loc[sc['ptype'] == 'h', 'formatted'] = \
                    sc.loc[sc['ptype'] == 'h', 'h_inout'].apply(lambda x: str(x) + '. ') \
                +   sc.loc[sc['ptype'] == 'h', 'h_title'] \
                +   sc.loc[sc['ptype'] == 'h', 'h_time'].apply(lambda x: ' - ' + str(x))
            
            sc.loc[sc['ptype'] == 'h', 'formatted'] = sc.loc[sc['ptype'] == 'h', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Scene Heading"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'            
                )
            return sc
        
        def action(self, sc: pd.DataFrame) -> pd.DataFrame:
            sc.loc[sc['ptype'] == 'a', 'formatted'] = \
                sc.loc[sc['ptype'] == 'a', 'pcontent'].apply(lambda x: str(x).lstrip('\s').rstrip('\s').replace('&', '&amp;'))
            sc.loc[sc['ptype'] == 'a', 'formatted'] = sc.loc[sc['ptype'] == 'a', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Action"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
            return sc
    
        def dialog(self, sc: pd.DataFrame) -> pd.DataFrame:
            
            #
            sc.loc[sc['ptype'] == 'd', 'formatted'] = \
                  sc.loc[sc['ptype'] == 'd', 'd_character'] \
                + sc.loc[sc['ptype'] == 'd', 'd_character_parenthesis'].apply(lambda x: '(' + str(x) + ')' if not pd.isnull(x) else '')
            
            sc.loc[sc['ptype'] == 'd', 'formatted'] = sc.loc[sc['ptype'] == 'd', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Character"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
                 
            sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'formatted'] \
                + sc.loc[(sc['ptype'] == 'd') & (sc['d_dialog_parenthesis'].isna()), 'd_dialog'].apply(
                    lambda x: '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                    )
            
            sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna()), 'formatted'] \
                + sc.loc[(sc['ptype'] == 'd') & (~sc['d_dialog_parenthesis'].isna())].agg(self.format_d_dialog, axis=1)
                            
            return sc
        
        def format_d_dialog(self, x: pd.DataFrame) -> str:
            '''
            formats dialog for XML output, to be used as a function for pandas agg or apply.
            
            Parameters
            ----------
            x : one row or column of pd.DataFrame depending on axis,
                to be used with the pandas agg or apply function
                DESCRIPTION.
    
            Returns
            -------
            str
                DESCRIPTION.
    
            '''
            pattern = r'[（(].*?[）)]'
            d_split = x['d_dialog'].split(pattern)
            f_dialog = ''
            for i, d in enumerate(d_split):
                ddialog = '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>' \
                    + str(d) \
                    + '</text>\n\t\t</para>\n'
                if i == 0:
                    f_dialog = f_dialog + ddialog
                else:
                    dparenthetical = '\t\t<para>\n\t\t\t<style basestylename="Parenthetical"/>\n\t\t\t<text>' \
                        + str(x['d_dialog_parenthesis'][i-1]) \
                        + '</text>\n\t\t</para>\n' 
                    f_dialog = f_dialog + dparenthetical + ddialog
            return f_dialog

######
class Translate(object):
    
    def __init__(self, sc: str = None):   
        super(Screenplay, self).__init__()
        
    def Baidu(str_to_trans='apple', lang_from='zh', lang_to='en'):
        time.sleep(2)
        #path_translink = 'http://api.fanyi.baidu.com/api/trans/vip/translate'
        path_translink = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
        httpClient = None
        
        appid = '20200218000385369'
        passcode = 'phecLsPI8EnkvjRHQ8HU'
        salt = randint(1e9, 9e9)
        q = str_to_trans
        for_sign = appid + q + str(salt) + passcode
        sign = hashlib.md5(for_sign.encode()).hexdigest()
        
        link_query = (path_translink + '?'
                      + 'appid=' + appid +'&'
                      + 'q=' + urllib.parse.quote(q) + '&'
                      + 'from=' + lang_from + '&'
                      + 'to=' + lang_to + '&'
                      + 'salt=' + str(salt) + '&'
                      + 'sign=' + sign)
        
        try:
            httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
            httpClient.request('GET', link_query)
    
            # response是HTTPResponse对象
            response = httpClient.getresponse()
            result_all = response.read().decode("utf-8")
            result = json.loads(result_all)
            print(result['trans_result'][0]['dst'])
            return result['trans_result'][0]['dst']  # str
    
        except Exception as e:
            print(e)
        finally:
            if httpClient:
               httpClient.close() 