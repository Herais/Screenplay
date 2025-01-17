# -*- coding: utf-8 -*-
"""
Created on Mon Feb 17 00:31:27 2020

@author: VXhpUS
"""
from pathlib import Path
import numpy as np
import pandas as pd
import re
from docx import Document
import time
from xml.dom import minidom
import lxml
import lxml.etree as et
import codecs
import math
import PyPDF2

## For Translate
import hashlib
from random import randint
import http.client
import urllib
import json

#%% Class Screenplay
class Screenplay(object):
    
    def __init__(self, sc: str = None):   

        super(Screenplay, self).__init__()
        self.sc = sc
        self.export = Export().to_OpenXML()
        self.groupSH2Name = {
            'A': 'Action Group',
            'D': 'Dialogue Group',
            'SH': 'Scene Heading Group'
            }
        self.element2basestyle = {
            'text':'Normal Text',
            'h': 'Scene Heading',
            'a': 'Action',
            'c': 'Character',
            'p': 'Parenthetical',
            'd': 'Dialogue',
            't': 'Transition',
            's': 'Shot',
            'uc': 'Unspoken Character'
            }
        self.translate = Translate()
        self.read = Read()
        self.parse = Parse()
        
        # Scene heading element
    

class Read(object):

    def __init__(self):
        
        super(Read, self).__init__()

    def auto(self, filepath: str) -> pd.DataFrame:
        fp = Path(filepath)
        
        try:
            filepath = fp.resolve()
        except FileNotFoundError:
            print('file does not exist')
        
        # get file extension
        filepath = str(filepath)
        fextention = str(filepath).split('.')[-1]
        
        if fextention in ['txt', 'text']:
            dfsc = self.text(filepath)         
        elif fextention in ['doc', 'docx']:
            dfsc = self.docx(filepath)            
        elif fextention in ['xml']:
            dfsc = self.openxml(filepath)       
        else:
            dfsc = None
            
        return dfsc
    
    @staticmethod
    def pdf(filepath: str,
                pat_sh=None,
                pat_d=None) -> pd.DataFrame:
        # creating an object 
        file = open(filepath, 'rb')
        pdfReader = PyPDF2.PdfFileReader(file)
        numPages = pdfReader.numPages
        txt = ''
        for n in range(numPages):
            txt += pdfReader.getPage(n).extractText()
        
        pdfReader.close()
        file.close()
                
    @staticmethod
    def openxml(filepath: str,
                pat_sh=None,
                pat_d=None,
                pat_shot=['FADE', 'CUT', 'DISSOLVE', 'INTERCUT'],
                xpath=None) -> pd.DataFrame:
        
        dfsc = pd.read_xml(filepath, xpath='paragraphs/para')
        dfsc['style'] = pd.read_xml(
            filepath, xpath='paragraphs/para/style')['basestyle']
        dfsc = dfsc.rename(columns={'style':'Type', 'text': 'Element'})
        dfsc['Type'] = dfsc['Type'].apply(str)
        dfsc['Element'] = dfsc['Element'].apply(str)

        dfsc.loc[dfsc['Type'] == 'Parenthetical', 'Type'] = 'Parenthetical_D'
        # Define Columns
        dfsc['Grp'] = None
        dfsc['Scene'] = None
        
        # Assign Grp D and A
        dfsc.loc[dfsc['Type'].str.contains('Dialogue|Character|Parenthetical_D'),
                 'Grp'] = 'D'
        dfsc.loc[dfsc['Type'].str.contains('Action'),
                 'Grp'] = 'A'

        # Identify SHOT and Transition in A Group
        idx_shot_and_transition = dfsc[(dfsc['Grp'] == 'A') & 
                                       dfsc['Element'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Grp'] = 'S'
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Type'] = 'Shot'
        
        pat_shot = ['FADE', 'CUT', 'DISSOLVE', 'INTERCUT']
        idx_transition = dfsc[(dfsc.index.isin(idx_shot_and_transition)) & 
            dfsc['Element'].str.contains('|'.join(pat_shot), flags=re.IGNORECASE)].index
        dfsc.loc[dfsc.index.isin(idx_transition), 'Grp'] = 'T'
        dfsc.loc[dfsc.index.isin(idx_transition), 'Type'] = 'Transition'

        
        # regenerate Scene Numbers
        idx_sh = dfsc[dfsc['Type']== 'Scene Heading'].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        nscenes = len(idx_sh)
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = list(range(1, nscenes+1))
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
        dfsc['Scene'].astype('int')
       

        # Clean na Element
        dfsc.dropna(subset=['Element'], inplace=True)
        dfsc.dropna(how='all')        
        dfsc = dfsc[['Scene', 'Element', 'Grp', 'Type']].copy()
        
        return dfsc
    
    @staticmethod
    def text(filepath: str,
             pat_sh=['INT./EXT.', 'INT/EXT', 'EXT', 'EXT\.', 'INT', 'INT\.', 
                      '内./外.', '内/外.', '内景', '外景', 
                      '内\.', '内,', '外\.', '外,',
                     ],
             pat_shot=['FADE', 'CUT', 'DISSOLVE', 'INTERCUT'],
             pat_d=None) -> pd.DataFrame:
        """
        This functions opens screennplays in text, assumming it somewhat
        conforms to the screenplay format.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        pat_sh : TYPE, optional
            DESCRIPTION. The default is None.
        pat_d : TYPE, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        dfsc : TYPE
            DESCRIPTION.

        """
        # Read text file
        f = open(filepath)
        lines = f.readlines()
        f.close()
        
        dfsc = pd.DataFrame(lines)
        dfsc.columns = ['raw']
        
        # Define Columns
        dfsc['Grp'] = None
        dfsc['Type'] = None
        dfsc['Scene'] = None
        
        # identify Scene Headings
        idx_sh = dfsc[dfsc['raw'].str.contains('|'.join(pat_sh))].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        dfsc.loc[dfsc.index.isin(idx_sh), 'Type'] = 'Scene Heading'
        
        # regenerate Scene Numbers
        nscenes = len(idx_sh)
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = list(range(1, nscenes+1))
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
        dfsc['Scene'].astype('int')
        # Remove leading and trailing spaces in Scene Heading
        dfsc.loc[dfsc['Grp'] == 'H', 'raw'] = \
            dfsc.loc[dfsc['Grp'] == 'H', 'raw'].apply(str.strip)
            
        # Identify Action
        dfsc['nspaces'] = dfsc['raw'].apply(lambda x: len(x)-len(x.lstrip()))
        mid = (dfsc['nspaces'].max() - dfsc['nspaces'].min()) //2 - 2
        dfsc.loc[(dfsc['Grp'] != 'H') & (dfsc['nspaces'] <= mid), 'Grp'] = 'A'
        dfsc.loc[dfsc['Grp'] == 'A', 'Type'] = 'Action'
        # Identify Dialogue
        dfsc.loc[(dfsc['Grp'] != 'H') & (dfsc['nspaces'] > mid), 'Grp'] = 'D'
        dfsc['raw'] = dfsc['raw'].apply(str.strip)
        
        # Identify Dchar in D Group
        idx_dchar = dfsc[(dfsc['Grp'] == 'D') & dfsc['raw'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_dchar), 'Type'] = 'Character'
        # Identify Dialog in D Group
        dfsc.loc[dfsc['Grp'] == 'D', 'Type'] = \
            dfsc.loc[dfsc['Grp'] == 'D', 'Type'].fillna('Dialogue')
            
        # Identify SHOT and Transition in A Group
        idx_shot_and_transition = dfsc[(dfsc['Grp'] == 'A') & dfsc['raw'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Grp'] = 'S'
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Type'] = 'Shot'
        
        idx_transition = dfsc[(dfsc.index.isin(idx_shot_and_transition)) & 
            dfsc['raw'].str.contains('|'.join(pat_shot), flags=re.IGNORECASE)].index
        dfsc.loc[dfsc.index.isin(idx_transition), 'Grp'] = 'T'
        dfsc.loc[dfsc.index.isin(idx_transition), 'Type'] = 'Transition'
        
        # Combine Elements
        dfsc['nelement'] = None
        nelement = dfsc[dfsc['raw'] == ''].shape[0]
        dfsc.loc[(dfsc['raw'] == ''), 'nelement'] = list(range(nelement))
        dfsc['nelement'].fillna(method='bfill', inplace=True)
 
        # Group to Elements
        dfsc = dfsc.groupby(['Scene', 'nelement', 'Grp', 'Type'])['raw'].apply(
            lambda x: ' '.join(x)).reset_index()
        dfsc = dfsc[dfsc['raw'] != '']
        dfsc.rename(columns={'raw': 'Element'}, inplace=True)
 
        # Clean SC
        dfsc['Element'] = dfsc['Element'].apply(lambda x: re.sub(':SC:', '', x).strip())
        
        # Organize Return
        dfsc = dfsc[['Scene', 'Element', 'Grp', 'Type']].copy()
        
        return dfsc.copy()
    
    @staticmethod
    def docx(filepath: str, 
            pat_sh=['INT./EXT.', 'INT/EXT', 'INT./EXT', 'INT/EXT.',
                    'EXT./INT.', 'EXT/INT', 'EXT./INT', 'EXT/INT.',
                    'EXT', 'EXT\.', 'INT', 'INT\.', 
                    '内./外.', '内/外.', '内景', '外景', 
                    '内\.', '内,', '外\.', '外,',
                   ],
             pat_shot=['FADE', 'CUT', 'DISSOLVE', 'INTERCUT'],
             pat_d=None,
             lang='en'
             ) -> pd.DataFrame:
        '''
        This function reads word docx into a screenplay structured pd.DataFrame, 
        oneline per row.

        Parameters
        ----------
        filepath : str
            DESCRIPTION.
        sctype : str, optional
            DESCRIPTION. The default is None.

        Returns
        -------
        script : pd.DataFrame 
            DESCRIPTION.

        '''
        
        # Read Docx
        document = Document(filepath)
        num_p = len(document.paragraphs)
        dfsc = []
        for num in range(num_p):
            dfsc.append(document.paragraphs[num].text)
        dfsc = pd.DataFrame(dfsc)
        dfsc.columns = ['Element']

        # Define Columns
        dfsc['Scene'] = None
        dfsc['Grp'] = None
        dfsc['Type'] = None
        
        # identify Scene Headings
        idx_sh = dfsc[dfsc['Element'].str.contains('|'.join(pat_sh))].index
        dfsc.loc[dfsc.index.isin(idx_sh), 'Grp'] = 'H'
        dfsc.loc[dfsc.index.isin(idx_sh), 'Type'] = 'Scene Heading'
        
        # Regenerate Scene Numbers
        dfsc.loc[dfsc.index.isin(idx_sh), 'Scene'] = [i 
            for i in range(1, len(idx_sh)+1)]
        dfsc['Scene'].fillna(method='ffill', inplace=True)
        dfsc['Scene'].fillna(-1, inplace=True)
            
        # Identify Dialogue
        if not pat_d:
            pat_d = '[:：]'
        idx_d = dfsc[dfsc['Element'].str.contains(pat_d)].index
        dfsc.loc[idx_d, 'Grp'] = 'D'
        dfsc.loc[idx_d, 'Type'] = 'Dialogue'
        
        # Identify Action
        dfsc['Grp'] = dfsc['Grp'].fillna('A')
        dfsc['Type'] = dfsc['Type'].fillna('Action')
 
        # Identify SHOT in A Group
        idx_shot_and_transition = dfsc[(dfsc['Grp'] == 'A') & dfsc['Element'].str.isupper()].index
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Grp'] = 'S'
        dfsc.loc[dfsc.index.isin(idx_shot_and_transition), 'Type'] = 'Shot'
        
        #Identify TRansition in A Group
        idx_transition = dfsc[(dfsc.index.isin(idx_shot_and_transition)) & 
                dfsc['Element'].str.contains('|'.join(pat_shot), flags=re.IGNORECASE)].index
        dfsc.loc[dfsc.index.isin(idx_transition), 'Grp'] = 'T'
        dfsc.loc[dfsc.index.isin(idx_transition), 'Type'] = 'Transition'
        
        # Clean Element with na
        dfsc['Element'] = dfsc['Element'].apply(lambda x: str(x).strip())
        dfsc = dfsc[dfsc['Element'] != ''].reset_index(drop=True)


        return dfsc
            
    @staticmethod
    def pdf(filepath: str,
                pat_sh=None,
                pat_d=None) -> pd.DataFrame:    
        pass

    @staticmethod
    def extract_location(x, 
                         pat_location:str ='[-——\.,]+'):
        if x['Element']:
            location = str(x['Element'])
        else: return ''
        if x['IE']:
            location = re.sub(re.escape(str(x['IE'])), '', location)
        if x['Time']:
            location = re.sub(re.escape(str(x['Time'])), '', location)
            
        location = re.sub(pat_location, '', location)
        return location.strip()




class Parse(object):
    
    def __init__(self):
        super(Parse, self).__init__()    
    
    @staticmethod
    def Scene_Heading(df: pd.DataFrame,
                 pat_sh=['INT./EXT.', 'INT/EXT', 'INT./EXT', 'INT/EXT.',
                         'EXT./INT.', 'EXT/INT', 'EXT./INT', 'EXT/INT.',
                         'EXT', 'EXT\.', 'INT', 'INT\.', 
                         '内./外.', '内/外.', '内景', '外景', 
                         '内\.', '内,', '外\.', '外,',
                        ],
                 pat_time:str = '[-——]+\s*(.*)',
                 pat_location:str = '[-——\.,]+'
                 ) -> pd.DataFrame:
        
        dfsc = df.copy()
        
        # Define Columns if not exist
        if 'IE' not in dfsc.columns:
            dfsc['IE'] = None
        if 'Location' not in dfsc.columns:
            dfsc['Location'] = None      
        if 'Time' not in dfsc.columns:
            dfsc['Time'] = None          
            
            
        # Break down Scene Headings
        idx_sh = dfsc[dfsc['Grp'] == 'H'].index
        
        # Extract Location
        dfsc.loc[dfsc.index.isin(idx_sh), 'IE'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), 'Element'].str.extract(
                '({})'.format('|'.join(pat_sh)), expand=False)
        
        # Extract Time
        dfsc.loc[dfsc.index.isin(idx_sh), 'Time'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), 'Element'].str.extract(
                pat_time, expand=False)
        
        # Extract Location
        dfsc.loc[dfsc.index.isin(idx_sh), 'Location'] = \
            dfsc.loc[dfsc.index.isin(idx_sh), :].apply(
                lambda x: Read.extract_location(x, pat_location=pat_location), 
                axis=1)
            
        return dfsc.copy()
    
    @staticmethod
    def D_Character_and_Dialogue(df: pd.DataFrame,
                 pat_dchar='[:：]',
                 pat_ddialogue=None,
                 pat_parenthetical='([(（].*[）)])',
                 ) -> pd.DataFrame():
        
        dfsc = df.copy()        
        dfD = dfsc.loc[dfsc['Grp'] == 'D', 'Element']
        
        #Extract Character and Dialogue in Grp D
        dfD_expanded = dfD.str.split(pat_dchar, n=1, expand=True)
        dfD_expanded.columns = ['Character', 'Dialogue']           
        dfD_melted = dfD_expanded.melt(ignore_index=False)
        dfD_melted.columns = ['Type_tmp', 'Element_tmp']
        dfsc_merged = dfsc.merge(dfD_melted, left_index=True, right_index=True, how='left')           
        dfsc_merged['Type'].update(dfsc_merged['Type_tmp'])
        dfsc_merged['Element'].update(dfsc_merged['Element_tmp'])
        dfsc_merged.drop(['Type_tmp', 'Element_tmp'], axis=1, inplace=True)
        dfsc = dfsc_merged.reset_index(drop=True)
            
        return dfsc.copy()
    
    @staticmethod
    def D_Character_Parenthetical(df: pd.DataFrame,
                                  pat_parenthetical='([(（].*[）)])',
                                  ) -> pd.DataFrame():
        
        dfsc = df.copy()
        dfC = dfsc.loc[dfsc['Type'] == 'Character']

        dfC_expanded = dfC['Element'].str.split(pat_parenthetical, expand=True)
        dfC_expanded = dfC_expanded.rename(columns= 
            {0:'Character', 1:'Parenthetical_C'})
        dfC_melted = dfC_expanded[['Character', 'Parenthetical_C']].melt(ignore_index=False)
        dfC_melted = dfC_melted.dropna(subset=['value'])
        dfC_melted['value'] = dfC_melted['value'].apply(str.strip)
        dfsc_merged = dfsc.merge(dfC_melted, left_index=True, 
                                 right_index=True, how='left')           
        dfsc_merged['Type'].update(dfsc_merged['variable'])
        dfsc_merged['Element'].update(dfsc_merged['value'])
        dfsc_merged.drop(['variable', 'value'], axis=1, inplace=True)     
        dfsc = dfsc_merged.reset_index(drop=True)
        dfsc = dfsc[dfsc['Element'] != '']
        
        dfsc.loc[dfsc.Type == 'Character', 'Element']
        
        return dfsc.copy()
                    
    @staticmethod   
    def D_Dialogue_Parenthetical(df: pd.DataFrame,
                             pat_parenthetical='([(（].*[）)])',
                            ) -> pd.DataFrame():
        dfsc = df.copy()
        dfD = dfsc.loc[dfsc['Type'] == 'Dialogue']
        dfD_expanded = dfD['Element'].str.split(pat_parenthetical, expand=True)
        dfD_melted = dfD_expanded.melt(ignore_index=False)
        dfD_melted.dropna(subset=['value'], inplace=True)
        dfD_melted = dfD_melted[dfD_melted['value'] != '']
        
        #Assign Label
        dfD_melted.loc[dfD_melted['value'].str.contains(pat_parenthetical), 
               'variable'] = 'Parenthetical_D'
        dfD_melted.loc[~dfD_melted['value'].str.contains(pat_parenthetical), 
               'variable'] = 'Dialogue'
        
        df_merged = dfsc.merge(dfD_melted, 
                                left_index=True, right_index=True, how='left')
        df_merged['Type'].update(df_merged['variable'])
        df_merged['Element'].update(df_merged['value']) 
        
        dfsc = df_merged[['Scene', 'Element', 'Grp', 'Type']].reset_index(
            drop=True)
        
        return dfsc.copy()
    
   
######################################################
class Export(object):
    
    def __init__(self):
        super(Export, self).__init__()
    
    class to_OpenXML(object):
        
        def __init__(self):
            pass
    
        def to_openxml(self, dfsc: pd.DataFrame, save: bool = False, file_path: str  = None) -> str:
            '''
            
    
            Parameters
            ----------
            sc : pd.DataFrame
                DESCRIPTION.
            save : str, optional
                DESCRIPTION. The default is None.
            file_path : str, optional
                DESCRIPTION. The default is None.
    
            Returns
            -------
            str
                DESCRIPTION.
    
            '''
            df = dfsc.copy()
            df['formatted'] = None
    
            header_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
            header_document = '<document type="Open Screenplay Format document" version="30">\n'
            document_info = '\t<info>\n\t</info>\n'            
            header_format = '\t<styles>\n\t</styles>\n'
            header_paragraphs = '\t<paragraphs>\n'
            
            footer_paragraphs = '\t</paragraphs>\n'
            section_titlepage = '\t<titlepage>\n\t</titlepage>\n'
            section_lists = '\t<lists>\n\t</lists>\n'
            footer_document = '</document>\n'
                
            df = self.heading(df)
            df = self.action(df)
            df = self.dialog(df)
    
            formatted_output = header_xml \
                             + header_document \
                             + document_info \
                             + header_format \
                             + header_paragraphs \
                             + sc['formatted'].sum() \
                             + footer_paragraphs \
                             + section_titlepage \
                             + section_lists \
                             + footer_document \
                                 
            if save:
                with codecs.open(file_path, "w", 'utf-8-sig') as f:
                    print(formatted_output, file=f)
                f.close()
                
            return formatted_output
        
        
        def heading(self, dfsc: pd.DataFrame) -> pd.DataFrame:
            df = dfsc.copy()
            df.loc[df['Grp'] == 'H', 'formatted'] = \
                    df.loc[df['Grp'] == 'H', 'IE'].apply(lambda x: str(x) + '. ') \
                +   df.loc[df['Grp'] == 'H', 'Location'] \
                +   df.loc[df['Grp'] == 'H', 'Time'].apply(
                    lambda x: ' - ' + str(x) if x else '')
            
            df.loc[df['Grp'] == 'H', 'formatted'] = \
                df.loc[df['Grp'] == 'H', 'Element'].apply(
                    lambda x: '\t\t<para>\n\t\t\t<style basestylename="Scene Heading"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'            
                )
            return df
        
        def action(self, dfsc: pd.DataFrame) -> pd.DataFrame:
            df = dfsc.copy()
            df.loc[df['Grp'] == 'A', 'formatted'] = \
                df.loc[df['Grp'] == 'A', 'Element'].apply(
                    lambda x: str(x).lstrip('\s').rstrip('\s').replace(
                        '&', '&amp;'))
            df.loc[df[df['Grp'] == 'A', 'formatted'] = \
                df.loc[df['ptype'] == 'a', 'Element'].apply(
                    lambda x: '\t\t<para>\n\t\t\t<style basestylename="Action"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
            return df.copy()
    
        def dialog(self, dfsc: pd.DataFrame) -> pd.DataFrame:
            
            df = dfsc.copy()
            df.loc[df['Type'] == 'Character', 'formatted'] = \
                  df['Type'] == 'Character', 'Element'] \
                + df['Type'] == 'Character', 'd_character_parenthesis'].apply(lambda x: '(' + str(x) + ')' if not pd.isnull(x) else '')
            
            df.loc[df['ptype'] == 'd', 'formatted'] = df.loc[df['ptype'] == 'd', 'formatted'].apply(
                lambda x: '\t\t<para>\n\t\t\t<style basestylename="Character"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                )
                 
            df.loc[(df['ptype'] == 'd') & (df['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  df.loc[(df['ptype'] == 'd') & (df['d_dialog_parenthesis'].isna()), 'formatted'] \
                + df.loc[(df['ptype'] == 'd') & (df['d_dialog_parenthesis'].isna()), 'd_dialog'].apply(
                    lambda x: '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>'
                        + str(x)
                        + '</text>\n\t\t</para>\n'
                    )
            
            df.loc[(df['ptype'] == 'd') & (~df['d_dialog_parenthesis'].isna()), 'formatted'] = \
                  df.loc[(df['ptype'] == 'd') & (~df['d_dialog_parenthesis'].isna()), 'formatted'] \
                + df.loc[(df['ptype'] == 'd') & (~df['d_dialog_parenthesis'].isna())].agg(self.format_d_dialog, axis=1)
                            
            return df
        
        def format_d_dialog(self, x: pd.DataFrame) -> str:
            '''
            formats dialog for XML output, to be used as a function for pandas agg or apply.
            
            Parameters
            ----------
            x : one row or column of pd.DataFrame depending on axis,
                to be used with the pandas agg or apply function
                DEdfRIPTION.
    
            Returns
            -------
            str
                DEdfRIPTION.
    
            '''
            pattern = r'[（(].*?[）)]'
            d_split = x['d_dialog'].split(pattern)
            f_dialog = ''
            for i, d in enumerate(d_split):
                ddialog = '\t\t<para>\n\t\t\t<style basestylename="Dialogue"/>\n\t\t\t<text>' \
                    + str(d) \
                    + '</text>\n\t\t</para>\n'
                if i == 0:
                    f_dialog = f_dialog + ddialog
                else:
                    dparenthetical = '\t\t<para>\n\t\t\t<style basestylename="Parenthetical"/>\n\t\t\t<text>' \
                        + str(x['d_dialog_parenthesis'][i-1]) \
                        + '</text>\n\t\t</para>\n' 
                    f_dialog = f_dialog + dparenthetical + ddialog
            return f_dialog

######
class Translate(object):
    
    def __init__(self, sc: str = None):   
        super(Translate, self).__init__()
    
    
    @staticmethod    
    def Baidu(s:str='apple', lang_from='zh', lang_to='en'):
        time.sleep(2)
        #path_translink = 'http://api.fanyi.baidu.com/api/trans/vip/translate'
        path_translink = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
        httpClient = None
        
        appid = '20200218000385369'
        passcode = 'phecLsPI8EnkvjRHQ8HU'
        salt = randint(1e9, 9e9)
        q = s
        for_sign = appid + q + str(salt) + passcode
        sign = hashlib.md5(for_sign.encode()).hexdigest()
        
        link_query = (path_translink + '?'
                      + 'appid=' + appid +'&'
                      + 'q=' + urllib.parse.quote(q) + '&'
                      + 'from=' + lang_from + '&'
                      + 'to=' + lang_to + '&'
                      + 'salt=' + str(salt) + '&'
                      + 'sign=' + sign)
        
        try:
            httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
            httpClient.request('GET', link_query)
    
            # response是HTTPResponse对象
            response = httpClient.getresponse()
            result_all = response.read().decode("utf-8")
            result = json.loads(result_all)
            print(result['trans_result'][0]['dst'])
            return result['trans_result'][0]['dst']  # str
    
        except Exception as e:
            print(e)
        finally:
            if httpClient:
               httpClient.close() 